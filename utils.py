import os
import requests
import streamlit as st
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


# Secure API Setup
GROQ_API_KEY = st.secrets["GROQ_API_KEY"]
GROQ_URL = "https://api.groq.com/openai/v1/chat/completions"
MODEL_NAME = "llama3-70b-8192"

headers = {
    "Authorization": f"Bearer {GROQ_API_KEY}",
    "Content-Type": "application/json"
}


def build_prompt(template, name, email, phone, job_title, company, experience, skills, education, linkedin_url):
    skills_list_bullet = '\n- ' + '\n- '.join(skill.strip() for skill in skills.split(','))
    skills_inline = ', '.join(skill.strip() for skill in skills.split(','))

    if template == "📋 Structured Pro":
        return f"""
# {name}

---

📩 **Email**: {email}  
📞 **Phone**: {phone}  
{f'🔗 **LinkedIn**: {linkedin_url}' if linkedin_url else ''}

---

### 🎯 Objective  
To secure a position as a **{job_title}** at **{company}**, leveraging my analytical skills to drive actionable insights.

---

### 🛠️ Technical Skills  
{skills_list_bullet}

---

### 🎓 Education  
{education}

---

### 💼 Projects & Experience  
{experience}
"""
    elif template == "🧘 Focused Minimal":
        return f"""
# {name}

---

**Email**: {email} | **Phone**: {phone}{' | **LinkedIn**: ' + linkedin_url if linkedin_url else ''}

---

**Objective**  
Seeking a position as {job_title} at {company} to apply analytical skills in a focused and impactful manner.

---

**Skills**  
{skills_inline}

---

**Education**  
{education}

---

**Experience**  
{experience}
"""
    else:
        return "Invalid template selected."



def generate_resume(name, email, phone, job_title, company, experience, skills, education, linkedin_url, template):
    try:
        prompt = build_prompt(template, name, email, phone, job_title, company, experience, skills, education, linkedin_url)
        prompt += "\n\nNote: Do NOT include any explanations, notes, or messages. Only return the final clean resume content."

        response = requests.post(GROQ_URL, headers=headers, json={
            "model": MODEL_NAME,
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.7
        })

        result = response.json()["choices"][0]["message"]["content"].strip()
        if "Note:" in result:
            result = result.split("Note:")[0].strip()

        return result
    except Exception as e:
        return f"❌ Error generating resume: {str(e)}"


def generate_cover_letter(name, email, phone, job_title, company, experience, skills, education):
    try:
        prompt = f"""
Write a formal and enthusiastic cover letter in Markdown format using these details:

Full Name: {name}  
Email: {email}  
Phone: {phone}  
Job Title: {job_title}  
Company: {company}  
Experience: {experience}  
Skills: {skills}  
Education: {education}

It should include a greeting, role interest, highlighted skills, and a positive closing.
"""
        response = requests.post(GROQ_URL, headers=headers, json={
            "model": MODEL_NAME,
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.7
        })

        result = response.json()["choices"][0]["message"]["content"].strip()
        if "Note:" in result:
            result = result.split("Note:")[0].strip()

        return result
    except Exception as e:
        return f"❌ Error generating cover letter: {str(e)}"



def convert_to_pdf(content_text, output_file="output.pdf"):
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    
    width, height = A4
    pdf.setFont("Helvetica", 12)

    # Starting position from top of the page
    x = 50
    y = height - 50
    line_height = 16

    lines = content_text.strip().split('\n')
    for line in lines:
        if y < 50:  # Add new page if space runs out
            pdf.showPage()
            pdf.setFont("Helvetica", 12)
            y = height - 50
        pdf.drawString(x, y, line)
        y -= line_height

    pdf.save()
    pdf_data = buffer.getvalue()
    buffer.close()
    return pdf_data



def convert_to_docx(content_md):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    lines = content_md.strip().split('\n')
    name_line = lines[0].strip().replace("#", "").strip()

    # Header
    title = doc.add_paragraph()
    title.alignment = 1  # Center align
    run = title.add_run(name_line)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 112, 192)  # Blue header

    doc.add_paragraph("\n")

    for line in lines[2:]:
        doc.add_paragraph(line)

    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io.read()
